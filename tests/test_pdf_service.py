"""Covers extract_pages' page_number contract (see its docstring in
pdf_service.py): real page numbers for formats that have them, None for
formats that don't — never a faked page 1 for docx/txt. Nothing here needs
the DB/auth fixtures from conftest.py; extract_pages only touches the
filesystem, so these write real temp files instead.
"""
import docx

from app.services import pdf_service


def test_docx_extraction_reports_no_page_number(tmp_path):
    path = tmp_path / "notes.docx"
    document = docx.Document()
    document.add_paragraph("First paragraph of the document.")
    document.add_paragraph("Second paragraph, still no real page concept.")
    document.save(path)

    pages = pdf_service.extract_pages(str(path), content_type="")

    assert len(pages) == 1
    page_number, text = pages[0]
    assert page_number is None
    assert "First paragraph" in text
    assert "Second paragraph" in text


def test_txt_extraction_reports_no_page_number(tmp_path):
    path = tmp_path / "notes.txt"
    path.write_text("Plain text has no pages either.", encoding="utf-8")

    pages = pdf_service.extract_pages(str(path), content_type="")

    assert len(pages) == 1
    page_number, text = pages[0]
    assert page_number is None
    assert text == "Plain text has no pages either."
